"""Create DOCX tool - works in sandbox or locally."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from pathlib import Path

from responses import CreateDocxResponse


async def create_docx(
    content: str,
    output_path: str,
    title: str | None = None,
    template_style: str = "normal",
) -> CreateDocxResponse:
    """
    Create DOCX file from markdown content.

    Args:
        content: Markdown content (supports # headings, paragraphs)
        output_path: Path to save DOCX file (e.g., "/workspace/report.docx")
        title: Optional document title
        template_style: Style template ("normal", "formal", "memo")

    Returns:
        CreateDocxResponse with output_path and file_size or error message

    Example:
        ```python
        result = await create_docx(
            content="# Title\n\nBody text",
            output_path="/workspace/report.docx",
            title="Report"
        )
        if result.success:
            print(f"Created: {result.output_path}, size: {result.file_size} bytes")
        ```
    """
    try:
        output_path_obj = Path(output_path)
        output_path_obj.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        # Set margins based on template style
        sections = doc.sections
        for section in sections:
            if template_style == "formal":
                section.top_margin = Inches(1.0)
                section.bottom_margin = Inches(1.0)
                section.left_margin = Inches(1.25)
                section.right_margin = Inches(1.25)
            elif template_style == "memo":
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(1.0)
                section.right_margin = Inches(1.0)
            else:
                section.top_margin = Inches(1.0)
                section.bottom_margin = Inches(1.0)
                section.left_margin = Inches(1.0)
                section.right_margin = Inches(1.0)

        if title:
            title_para = doc.add_paragraph(title)
            title_para.style = "Title"
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Parse markdown-like content
        for para_text in content.split("\n\n"):
            if para_text.strip():
                if para_text.startswith("# "):
                    para = doc.add_paragraph(para_text[2:])
                    para.style = "Heading 1"  # type: ignore[assignment]
                elif para_text.startswith("## "):
                    para = doc.add_paragraph(para_text[3:])
                    para.style = "Heading 2"  # type: ignore[assignment]
                elif para_text.startswith("### "):
                    para = doc.add_paragraph(para_text[4:])
                    para.style = "Heading 3"  # type: ignore[assignment]
                else:
                    para = doc.add_paragraph(para_text)
                    if template_style == "formal":
                        for run in para.runs:
                            run.font.size = Pt(12)
                    elif template_style == "memo":
                        for run in para.runs:
                            run.font.size = Pt(11)

        doc.save(str(output_path_obj))

        return CreateDocxResponse(
            success=True,
            output_path=str(output_path_obj.absolute()),
            file_size=output_path_obj.stat().st_size,
            error=None,
        )
    except Exception as e:
        return CreateDocxResponse(
            success=False, error=f"Error creating DOCX: {str(e)}", output_path=None, file_size=None
        )
